#!/usr/bin/env python3
"""Local, read-only credential candidate review for selected inventory files.

The module intentionally never serializes a discovered secret. Candidate output
contains only a presence flag, its length and a fixed mask.
"""

from __future__ import annotations

import argparse
import configparser
import csv
import hashlib
import ipaddress
import io
import json
import logging
import re
import sys
import unicodedata
import warnings
import zipfile
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Iterator, Mapping
from xml.etree import ElementTree


APP_VERSION = "0.28.4-beta.2"
ROOT_CLIENT_LABEL = "(radice)"
MAX_FILE_BYTES = 20 * 1024 * 1024
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_TEXT_CHARACTERS = 2_000_000
MAX_RECORDS_PER_FILE = 5_000
MAX_PDF_PAGES = 200
MAX_SECURE_REVIEW_STDIN_BYTES = 64 * 1024 * 1024
MAX_OFFICE_PASSWORD_CHARACTERS = 1_024
SOURCE_MAPPING_SCHEMA_VERSION = 1
SOURCE_MAPPING_FIELDS = ("title", "username", "secret", "uri")
MAX_SOURCE_MAPPING_PROFILE_BYTES = 16 * 1024
MAX_SOURCE_MAPPING_PROFILE_NAME_CHARACTERS = 80
MAX_SOURCE_MAPPING_ALIASES_PER_FIELD = 8
MAX_SOURCE_MAPPING_ALIAS_CHARACTERS = 80
OLE_COMPOUND_FILE_SIGNATURE = bytes.fromhex("D0CF11E0A1B11AE1")

# Third-party document readers can emit advisory warnings on stderr. The CLI
# stdout is a strict JSON contract consumed by Windows PowerShell 5.1, so known
# parser warnings are suppressed here and represented as controlled per-file
# warnings by analyze_files instead.
logging.getLogger("pypdf").setLevel(logging.ERROR)
warnings.filterwarnings("ignore", module=r"^(openpyxl|docx|pypdf)(\.|$)")

REVIEWABLE_EXTENSIONS = {
    ".txt",
    ".csv",
    ".tsv",
    ".json",
    ".xml",
    ".yaml",
    ".yml",
    ".ini",
    ".cfg",
    ".conf",
    ".env",
    ".properties",
    ".docx",
    ".xlsx",
    ".xls",
    ".pdf",
}


class ReviewError(RuntimeError):
    """A safe, user-facing review error."""


class ExcelPasswordRequired(ReviewError):
    """A modern Excel document is encrypted and needs a password."""


class ExcelPasswordRejected(ReviewError):
    """The supplied Excel document password was rejected."""


class ExcelEncryptionReaderUnavailable(ReviewError):
    """The optional in-memory Office decryption dependency is unavailable."""


class LegacyExcelConversionRequired(ReviewError):
    """A legacy XLS source must be converted before review."""


class UnsupportedReviewFormat(ReviewError):
    """The selected source format is outside the review contract."""


class ReviewFileTooLarge(ReviewError):
    """The selected source exceeds the bounded per-document limit."""


class ReviewSourceUnavailable(ReviewError):
    """The selected source is absent, linked or outside the trusted root."""


@dataclass(frozen=True)
class SourceMappingProfile:
    """Validated, secret-free mapping from source labels to candidate fields."""

    schema_version: int
    name: str
    fields: dict[str, tuple[str, ...]]
    digest: str

    def document(self) -> dict[str, object]:
        return {
            "schema_version": self.schema_version,
            "name": self.name,
            "fields": {
                field_name: list(self.fields[field_name])
                for field_name in SOURCE_MAPPING_FIELDS
            },
            "digest": self.digest,
        }


@dataclass(frozen=True)
class CredentialCandidate:
    candidate_id: str
    source_relative_path: str
    source_sha256: str
    client: str
    location: str
    title: str
    username: str
    uri: str
    secret_present: bool
    secret_mask: str
    secret_length: int
    status: str
    confidence: str
    source_password_required: bool = False
    fields_detected: list[str] = field(default_factory=list)
    source_mapping_digest: str = ""
    source_mapping_profile: dict[str, object] | None = None


@dataclass(frozen=True)
class ProtectedExcelIssue:
    relative_path: str
    status: str


@dataclass(frozen=True)
class ReviewIssueSummary:
    """One aggregate review issue without source identifiers."""

    reason_code: str
    extension: str
    count: int


@dataclass(frozen=True)
class ReviewResult:
    root: str
    reviewed_at_utc: str
    selected_files: int
    analyzed_files: int
    candidate_count: int
    ready_count: int
    incomplete_count: int
    candidates: list[CredentialCandidate] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    protected_excel_issues: list[ProtectedExcelIssue] = field(default_factory=list)
    issues: list[ReviewIssueSummary] = field(default_factory=list)
    source_mapping_profile_name: str = "Rilevamento automatico"
    source_mapping_profile_digest: str = ""


def _normalized_key(value: object) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    ascii_text = "".join(character for character in text if not unicodedata.combining(character))
    return re.sub(r"[^a-z0-9]", "", ascii_text.casefold())


def _canonical_source_mapping_document(
    name: str, fields: Mapping[str, tuple[str, ...]]
) -> dict[str, object]:
    return {
        "schema_version": SOURCE_MAPPING_SCHEMA_VERSION,
        "name": name,
        "fields": {
            field_name: list(fields[field_name]) for field_name in SOURCE_MAPPING_FIELDS
        },
    }


def normalize_source_mapping_profile(value: object) -> SourceMappingProfile | None:
    """Validate and canonicalize an optional exact-match source mapping profile."""

    if value is None:
        return None
    if not isinstance(value, Mapping):
        raise ReviewError("Il profilo di mappatura sorgente deve essere un oggetto JSON.")
    allowed_keys = {"schema_version", "name", "fields", "digest"}
    if set(value) - allowed_keys:
        raise ReviewError("Il profilo di mappatura contiene campi non riconosciuti.")
    if value.get("schema_version") != SOURCE_MAPPING_SCHEMA_VERSION:
        raise ReviewError("La versione del profilo di mappatura non è supportata.")
    name = str(value.get("name", "")).strip()
    if (
        not name
        or len(name) > MAX_SOURCE_MAPPING_PROFILE_NAME_CHARACTERS
        or re.search(r"[\x00-\x1f\x7f]", name)
    ):
        raise ReviewError("Il nome del profilo di mappatura non è valido.")
    raw_fields = value.get("fields")
    if not isinstance(raw_fields, Mapping) or set(raw_fields) != set(
        SOURCE_MAPPING_FIELDS
    ):
        raise ReviewError(
            "Il profilo deve dichiarare esattamente title, username, secret e uri."
        )

    normalized_fields: dict[str, tuple[str, ...]] = {}
    alias_owner: dict[str, str] = {}
    for field_name in SOURCE_MAPPING_FIELDS:
        raw_aliases = raw_fields.get(field_name)
        if (
            not isinstance(raw_aliases, list)
            or len(raw_aliases) > MAX_SOURCE_MAPPING_ALIASES_PER_FIELD
        ):
            raise ReviewError(
                f"La mappatura {field_name} deve essere una lista di massimo "
                f"{MAX_SOURCE_MAPPING_ALIASES_PER_FIELD} etichette."
            )
        normalized_aliases: list[str] = []
        for raw_alias in raw_aliases:
            if (
                not isinstance(raw_alias, str)
                or not raw_alias.strip()
                or len(raw_alias.strip()) > MAX_SOURCE_MAPPING_ALIAS_CHARACTERS
                or re.search(r"[\x00-\x1f\x7f]", raw_alias)
            ):
                raise ReviewError(
                    f"Una etichetta della mappatura {field_name} non è valida."
                )
            normalized = _normalized_key(raw_alias)
            if not normalized or normalized in normalized_aliases:
                raise ReviewError(
                    f"La mappatura {field_name} contiene etichette vuote o duplicate."
                )
            previous_owner = alias_owner.get(normalized)
            if previous_owner is not None:
                raise ReviewError(
                    "Una stessa etichetta sorgente non può alimentare due campi: "
                    f"{previous_owner} e {field_name}."
                )
            alias_owner[normalized] = field_name
            normalized_aliases.append(normalized)
        normalized_fields[field_name] = tuple(normalized_aliases)

    if not normalized_fields["secret"] or not (
        normalized_fields["username"] or normalized_fields["uri"]
    ):
        raise ReviewError(
            "Il profilo deve mappare la password e almeno uno fra username e URL/host."
        )

    canonical = _canonical_source_mapping_document(name, normalized_fields)
    digest = hashlib.sha256(
        json.dumps(
            canonical,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()
    supplied_digest = value.get("digest")
    if supplied_digest is not None and str(supplied_digest).strip().lower() != digest:
        raise ReviewError("Il digest del profilo di mappatura non corrisponde al contenuto.")
    return SourceMappingProfile(
        schema_version=SOURCE_MAPPING_SCHEMA_VERSION,
        name=name,
        fields=normalized_fields,
        digest=digest,
    )


TITLE_KEYS = {
    _normalized_key(value)
    for value in (
        "title",
        "titolo",
        "name",
        "nome",
        "resource",
        "risorsa",
        "service",
        "servizio",
        "application",
        "applicazione",
        "descrizione",
    )
}
USERNAME_KEYS = {
    _normalized_key(value)
    for value in (
        "username",
        "user",
        "utente",
        "login",
        "login name",
        "user name",
        "nome utente",
        "email",
        "e-mail",
        "account",
    )
}
SECRET_KEYS = {
    _normalized_key(value)
    for value in (
        "password",
        "pass",
        "passwd",
        "pwd",
        "secret",
        "segreto",
        "pin",
        "token",
        "api key",
        "apikey",
        "chiave api",
    )
}
URI_KEYS = {
    _normalized_key(value)
    for value in (
        "url",
        "uri",
        "website",
        "sito",
        "link",
        "host",
        "hostname",
        "server",
        "ip",
        "ip address",
        "server ip",
        "ip server",
        "ipv4",
        "ipv6",
        "indirizzo",
        "indirizzo ip",
        "indirizzo ipv4",
        "indirizzo ipv6",
    )
}
ALL_CREDENTIAL_KEYS = TITLE_KEYS | USERNAME_KEYS | SECRET_KEYS | URI_KEYS
USERNAME_PREFIX_KEYS = tuple(alias for alias in USERNAME_KEYS if len(alias) >= 4)
SECRET_PREFIX_KEYS = tuple(alias for alias in SECRET_KEYS if len(alias) >= 4)
URI_PREFIX_KEYS = tuple(alias for alias in URI_KEYS if len(alias) >= 4)


def _utc_now() -> str:
    return datetime.now(tz=timezone.utc).isoformat().replace("+00:00", "Z")


def _scalar(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, (str, int, float, bool)):
        return str(value).strip()
    return ""


def _matches_normalized_field_key(
    normalized: str, aliases: set[str], prefix_aliases: tuple[str, ...] = ()
) -> bool:
    if normalized in aliases:
        return True
    return normalized.endswith(prefix_aliases) if prefix_aliases else False


def _matches_field_key(key: object, aliases: set[str], allow_prefix: bool) -> bool:
    normalized = _normalized_key(key)
    if not allow_prefix:
        prefix_aliases = ()
    elif aliases is USERNAME_KEYS:
        prefix_aliases = USERNAME_PREFIX_KEYS
    elif aliases is SECRET_KEYS:
        prefix_aliases = SECRET_PREFIX_KEYS
    elif aliases is URI_KEYS:
        prefix_aliases = URI_PREFIX_KEYS
    else:
        prefix_aliases = tuple(alias for alias in aliases if len(alias) >= 4)
    # Environment/configuration files commonly use keys such as DB_PASSWORD
    # or PORTAL_USERNAME. Short aliases (for example "ip") are kept exact to
    # avoid accidental suffix matches in unrelated words.
    return _matches_normalized_field_key(normalized, aliases, prefix_aliases)


def _find_field(
    record: dict[object, object], aliases: set[str], *, allow_prefix: bool = False
) -> tuple[bool, str]:
    for key, value in record.items():
        if _matches_field_key(key, aliases, allow_prefix):
            return True, _scalar(value)
    return False, ""


def _candidate_fields(
    record: dict[object, object],
    source_mapping_profile: SourceMappingProfile | None = None,
) -> tuple[tuple[bool, str], tuple[bool, str], tuple[bool, str], tuple[bool, str]]:
    if source_mapping_profile is not None:
        values_by_key: dict[str, list[str]] = {}
        for key, value in record.items():
            normalized = _normalized_key(key)
            if normalized:
                values_by_key.setdefault(normalized, []).append(_scalar(value))

        mapped_fields: list[tuple[bool, str]] = []
        for field_name in SOURCE_MAPPING_FIELDS:
            selected: tuple[bool, str] = (False, "")
            matched_alias = ""
            for alias in source_mapping_profile.fields[field_name]:
                matches = values_by_key.get(alias, [])
                if len(matches) > 1:
                    raise ReviewError(
                        f"Il profilo {source_mapping_profile.name} trova più colonne "
                        f"equivalenti per {field_name}."
                    )
                if matches:
                    if matched_alias:
                        raise ReviewError(
                            f"Il profilo {source_mapping_profile.name} trova sia "
                            f"{matched_alias} sia {alias} per {field_name}."
                        )
                    matched_alias = alias
                    selected = (True, matches[0])
            mapped_fields.append(selected)
        return tuple(mapped_fields)  # type: ignore[return-value]

    title = (False, "")
    username = (False, "")
    secret = (False, "")
    uri = (False, "")
    for key, value in record.items():
        normalized = _normalized_key(key)
        scalar: str | None = None
        if not title[0] and _matches_normalized_field_key(normalized, TITLE_KEYS):
            scalar = _scalar(value)
            title = (True, scalar)
        if not username[0] and _matches_normalized_field_key(
            normalized, USERNAME_KEYS, USERNAME_PREFIX_KEYS
        ):
            scalar = _scalar(value) if scalar is None else scalar
            username = (True, scalar)
        if not secret[0] and _matches_normalized_field_key(
            normalized, SECRET_KEYS, SECRET_PREFIX_KEYS
        ):
            scalar = _scalar(value) if scalar is None else scalar
            secret = (True, scalar)
        if not uri[0] and _matches_normalized_field_key(
            normalized, URI_KEYS, URI_PREFIX_KEYS
        ):
            scalar = _scalar(value) if scalar is None else scalar
            uri = (True, scalar)
        if title[0] and username[0] and secret[0] and uri[0]:
            break
    return title, username, secret, uri


IPV4_CANDIDATE_PATTERN = re.compile(
    r"(?<![0-9])(?:[0-9]{1,3}\.){3}[0-9]{1,3}(?![0-9])"
)


def _ip_from_scalar(value: object) -> str:
    """Return the first valid IP address embedded in a scalar value."""

    text = _scalar(value)
    if not text:
        return ""
    for match in IPV4_CANDIDATE_PATTERN.finditer(text):
        try:
            return str(ipaddress.ip_address(match.group(0)))
        except ValueError:
            continue

    # IPv6 is deliberately considered only from compact tokens. This avoids
    # treating ordinary prose containing colons as a host address.
    for token in re.split(r"[\s,;]+", text):
        candidate = token.strip("()[]{}<>'\"")
        if candidate.count(":") < 2:
            continue
        if "/" in candidate:
            candidate = candidate.split("/", 1)[0]
        if "%" in candidate:
            candidate = candidate.split("%", 1)[0]
        try:
            address = ipaddress.ip_address(candidate)
        except ValueError:
            continue
        if address.version == 6:
            return str(address)
    return ""


def _find_ip_address(record: dict[object, object]) -> str:
    """Find an IP in non-secret fields when no explicit URL/host was found."""

    for key, value in record.items():
        if _matches_field_key(key, SECRET_KEYS, True):
            continue
        address = _ip_from_scalar(value)
        if address:
            return address
    return ""


def _has_credential_key(
    record: dict[object, object],
    source_mapping_profile: SourceMappingProfile | None = None,
) -> bool:
    if source_mapping_profile is not None:
        aliases = {
            alias
            for field_name in SOURCE_MAPPING_FIELDS
            for alias in source_mapping_profile.fields[field_name]
        }
        return any(_normalized_key(key) in aliases for key in record)
    return any(
        _matches_field_key(key, TITLE_KEYS, False)
        or _matches_field_key(key, USERNAME_KEYS, True)
        or _matches_field_key(key, SECRET_KEYS, True)
        or _matches_field_key(key, URI_KEYS, True)
        for key in record
    )


def _make_candidate(
    record: dict[object, object],
    *,
    relative_path: str,
    source_hash: str,
    client: str,
    location: str,
    source_password_required: bool = False,
    source_mapping_profile: SourceMappingProfile | None = None,
) -> CredentialCandidate | None:
    (
        (title_found, title),
        (username_found, username),
        (secret_found, secret),
        (uri_found, uri),
    ) = _candidate_fields(record, source_mapping_profile)
    if not uri and source_mapping_profile is None:
        detected_ip = _find_ip_address(record)
        if detected_ip:
            uri_found, uri = True, detected_ip

    secret_present = bool(secret_found and secret)
    if not secret_found and not (username_found and (title_found or uri_found)):
        return None

    if not title:
        title = Path(relative_path).stem
    fields_detected: list[str] = []
    if title_found:
        fields_detected.append("title")
    if username_found:
        fields_detected.append("username")
    if secret_found:
        fields_detected.append("secret")
    if uri_found:
        fields_detected.append("uri")

    ready = secret_present and bool(username or uri)
    confidence = "high" if ready and (title_found or uri_found) else "medium"
    if not secret_present:
        confidence = "low"
    identifier_parts = [relative_path, location, title, username, uri, source_hash]
    if source_mapping_profile is not None:
        identifier_parts.append(source_mapping_profile.digest)
    identifier_material = "\x1f".join(identifier_parts).encode(
        "utf-8", errors="surrogatepass"
    )
    candidate_id = hashlib.sha256(identifier_material).hexdigest()[:16]

    return CredentialCandidate(
        candidate_id=candidate_id,
        source_relative_path=relative_path,
        source_sha256=source_hash,
        client=client,
        location=location,
        title=title,
        username=username,
        uri=uri,
        secret_present=secret_present,
        secret_mask="********" if secret_present else "",
        secret_length=len(secret) if secret_present else 0,
        status="ready" if ready else "incomplete",
        confidence=confidence,
        source_password_required=source_password_required,
        fields_detected=fields_detected,
        source_mapping_digest=(
            source_mapping_profile.digest if source_mapping_profile is not None else ""
        ),
        source_mapping_profile=(
            source_mapping_profile.document()
            if source_mapping_profile is not None
            else None
        ),
    )


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "cp1252", "latin-1"):
        try:
            text = raw.decode(encoding)
            return text[:MAX_TEXT_CHARACTERS]
        except UnicodeDecodeError:
            continue
    raise ReviewError("Codifica testuale non riconosciuta.")


def _strip_wrapping_quotes(value: str) -> str:
    value = value.strip()
    if len(value) >= 2 and value[0] == value[-1] and value[0] in {"'", '"'}:
        return value[1:-1]
    return value


def _key_value_records(
    text: str,
    label: str = "testo",
    source_mapping_profile: SourceMappingProfile | None = None,
) -> Iterator[tuple[str, dict[str, str]]]:
    record: dict[str, str] = {}
    start_line = 1

    def emit(end_line: int) -> tuple[str, dict[str, str]] | None:
        nonlocal record, start_line
        if not record or not _has_credential_key(record, source_mapping_profile):
            record = {}
            return None
        location = f"{label}, righe {start_line}-{end_line}"
        result = (location, record)
        record = {}
        return result

    lines = text.splitlines()
    for line_number, original_line in enumerate(lines, start=1):
        line = original_line.strip()
        if not line:
            result = emit(line_number - 1)
            if result:
                yield result
            start_line = line_number + 1
            continue
        if line.startswith(("#", ";", "//")):
            continue
        line = re.sub(r"^[-*]\s+", "", line)
        match = re.match(r"^([^:=\t]{1,100})\s*[:=\t]\s*(.*)$", line)
        if not match:
            continue
        key = match.group(1).strip()
        value = _strip_wrapping_quotes(match.group(2))
        normalized = _normalized_key(key)
        if normalized in {_normalized_key(existing) for existing in record}:
            result = emit(line_number - 1)
            if result:
                yield result
            start_line = line_number
        record[key] = value
    result = emit(len(lines))
    if result:
        yield result


def _csv_records(text: str, delimiter: str) -> Iterator[tuple[str, dict[str, str]]]:
    reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
    if not reader.fieldnames:
        return
    for row_number, row in enumerate(reader, start=2):
        if row_number > MAX_RECORDS_PER_FILE + 1:
            break
        yield f"riga {row_number}", {str(key): _scalar(value) for key, value in row.items() if key}


def _json_records(
    text: str, source_mapping_profile: SourceMappingProfile | None = None
) -> Iterator[tuple[str, dict[object, object]]]:
    try:
        document = json.loads(text)
    except json.JSONDecodeError as exc:
        raise ReviewError(f"JSON non valido alla riga {exc.lineno}.") from exc

    emitted = 0

    def visit(value: object, path: str) -> Iterator[tuple[str, dict[object, object]]]:
        nonlocal emitted
        if emitted >= MAX_RECORDS_PER_FILE:
            return
        if isinstance(value, dict):
            if _has_credential_key(value, source_mapping_profile):
                emitted += 1
                yield path, value
            for key, child in value.items():
                if isinstance(child, (dict, list)):
                    safe_key = str(key).replace("~", "~0").replace("/", "~1")
                    yield from visit(child, f"{path}/{safe_key}")
        elif isinstance(value, list):
            for index, child in enumerate(value):
                if isinstance(child, (dict, list)):
                    yield from visit(child, f"{path}/{index}")

    yield from visit(document, "$")


def _xml_records(
    raw: bytes, source_mapping_profile: SourceMappingProfile | None = None
) -> Iterator[tuple[str, dict[str, str]]]:
    normalized = raw.upper()
    if b"<!DOCTYPE" in normalized or b"<!ENTITY" in normalized:
        raise ReviewError("XML con DTD o entità non consentite.")
    try:
        root = ElementTree.fromstring(raw)
    except ElementTree.ParseError as exc:
        raise ReviewError("XML non valido.") from exc

    emitted = 0
    for index, element in enumerate(root.iter(), start=1):
        children = list(element)
        if not children:
            continue
        record: dict[str, str] = {}
        for child in children:
            if not list(child):
                tag = child.tag.rsplit("}", 1)[-1]
                record[tag] = (child.text or "").strip()
        if record and _has_credential_key(record, source_mapping_profile):
            emitted += 1
            yield f"elemento {element.tag.rsplit('}', 1)[-1]} #{index}", record
            if emitted >= MAX_RECORDS_PER_FILE:
                break


def _ini_records(
    text: str, source_mapping_profile: SourceMappingProfile | None = None
) -> Iterator[tuple[str, dict[str, str]]]:
    parser = configparser.ConfigParser(interpolation=None, strict=False)
    parser.optionxform = str
    try:
        parser.read_string(text)
    except configparser.Error:
        yield from _key_value_records(
            text, source_mapping_profile=source_mapping_profile
        )
        return
    for section in parser.sections():
        yield f"sezione [{section}]", dict(parser.items(section))
    if parser.defaults():
        yield "sezione [DEFAULT]", dict(parser.defaults())


def _preflight_zip(source: Path | io.BytesIO) -> None:
    if isinstance(source, io.BytesIO):
        source.seek(0)
    try:
        with zipfile.ZipFile(source) as archive:
            total = sum(member.file_size for member in archive.infolist())
            if total > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
                raise ReviewError("Archivio Office troppo grande dopo la decompressione.")
            if any(member.flag_bits & 0x1 for member in archive.infolist()):
                raise ReviewError("Archivio Office cifrato non supportato.")
    except zipfile.BadZipFile as exc:
        raise ReviewError("Documento Office non valido o danneggiato.") from exc
    finally:
        if isinstance(source, io.BytesIO):
            source.seek(0)


def _is_ole_compound_file(path: Path) -> bool:
    try:
        with path.open("rb") as stream:
            return stream.read(len(OLE_COMPOUND_FILE_SIGNATURE)) == OLE_COMPOUND_FILE_SIGNATURE
    except OSError:
        return False


def _decrypt_xlsx_in_memory(path: Path, password: str | None) -> io.BytesIO:
    if password is not None and (
        not password or len(password) > MAX_OFFICE_PASSWORD_CHARACTERS
    ):
        raise ExcelPasswordRejected("La password del file Excel non è valida.")
    try:
        import msoffcrypto
        from msoffcrypto import exceptions as msoffcrypto_exceptions
    except ImportError as exc:
        raise ExcelEncryptionReaderUnavailable(
            "Supporto Excel cifrato non disponibile: installare msoffcrypto-tool 6.0.0."
        ) from exc

    decrypted = io.BytesIO()
    try:
        with path.open("rb") as encrypted_stream:
            office_file = msoffcrypto.OfficeFile(encrypted_stream)
            if not office_file.is_encrypted():
                raise ReviewError("Documento Excel OLE non cifrato o non riconosciuto.")
            if password is None:
                raise ExcelPasswordRequired("Il file Excel è protetto da password.")
            office_file.load_key(password=password, verify_password=True)
            office_file.decrypt(decrypted, verify_integrity=True)
        decrypted.seek(0)
        _preflight_zip(decrypted)
        return decrypted
    except msoffcrypto_exceptions.InvalidKeyError as exc:
        decrypted.close()
        raise ExcelPasswordRejected("Password del file Excel non corretta.") from exc
    except msoffcrypto_exceptions.DecryptionError as exc:
        decrypted.close()
        raise ExcelPasswordRejected("Password del file Excel non corretta o file danneggiato.") from exc
    except (msoffcrypto_exceptions.FileFormatError, msoffcrypto_exceptions.ParseError) as exc:
        decrypted.close()
        raise ReviewError("Formato del file Excel cifrato non riconosciuto.") from exc
    except Exception:
        decrypted.close()
        raise


def _xlsx_records(
    path: Path, password: str | None = None
) -> Iterator[tuple[str, dict[str, str]]]:
    decrypted: io.BytesIO | None = None
    workbook_source: Path | io.BytesIO = path
    if zipfile.is_zipfile(path):
        _preflight_zip(path)
    elif _is_ole_compound_file(path):
        decrypted = _decrypt_xlsx_in_memory(path, password)
        workbook_source = decrypted
    else:
        _preflight_zip(path)
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        if decrypted is not None:
            decrypted.close()
        raise ReviewError("Parser XLSX non disponibile.") from exc

    try:
        workbook = load_workbook(
            workbook_source, read_only=True, data_only=False, keep_links=False
        )
    except Exception:
        if decrypted is not None:
            decrypted.close()
        raise
    try:
        emitted = 0
        for worksheet in workbook.worksheets:
            headers: list[str] | None = None
            header_row = 0
            for row_number, values in enumerate(worksheet.iter_rows(values_only=True), start=1):
                if not any(value is not None and str(value).strip() for value in values):
                    continue
                if headers is None:
                    headers = [
                        _scalar(value) or f"colonna_{index}"
                        for index, value in enumerate(values, start=1)
                    ]
                    header_row = row_number
                    continue
                record = {
                    headers[index]: _scalar(value)
                    for index, value in enumerate(values)
                    if index < len(headers)
                }
                emitted += 1
                yield f"foglio {worksheet.title}, riga {row_number}", record
                if emitted >= MAX_RECORDS_PER_FILE:
                    return
            if headers is not None and header_row and emitted >= MAX_RECORDS_PER_FILE:
                return
    finally:
        workbook.close()
        if decrypted is not None:
            decrypted.close()


def _docx_records(
    path: Path, source_mapping_profile: SourceMappingProfile | None = None
) -> Iterator[tuple[str, dict[str, str]]]:
    _preflight_zip(path)
    try:
        from docx import Document
    except ImportError as exc:
        raise ReviewError("Parser DOCX non disponibile.") from exc

    document = Document(path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    yield from _key_value_records(
        paragraph_text, "paragrafi", source_mapping_profile
    )
    emitted = 0
    for table_number, table in enumerate(document.tables, start=1):
        if not table.rows:
            continue
        headers = [cell.text.strip() or f"colonna_{index}" for index, cell in enumerate(table.rows[0].cells, start=1)]
        for row_number, row in enumerate(table.rows[1:], start=2):
            record = {
                headers[index]: cell.text.strip()
                for index, cell in enumerate(row.cells)
                if index < len(headers)
            }
            emitted += 1
            yield f"tabella {table_number}, riga {row_number}", record
            if emitted >= MAX_RECORDS_PER_FILE:
                return


def _pdf_records(
    path: Path, source_mapping_profile: SourceMappingProfile | None = None
) -> Iterator[tuple[str, dict[str, str]]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ReviewError("Parser PDF non disponibile.") from exc

    reader = PdfReader(str(path), strict=False)
    if reader.is_encrypted:
        try:
            if reader.decrypt("") == 0:
                raise ReviewError("PDF cifrato: inserimento password non disponibile.")
        except ReviewError:
            raise
        except Exception as exc:
            raise ReviewError("PDF cifrato: inserimento password non disponibile.") from exc
    for page_number, page in enumerate(reader.pages[:MAX_PDF_PAGES], start=1):
        try:
            text = (page.extract_text() or "")[:MAX_TEXT_CHARACTERS]
        except Exception as exc:
            raise ReviewError(f"Testo PDF non leggibile a pagina {page_number}.") from exc
        yield from _key_value_records(
            text, f"pagina {page_number}", source_mapping_profile
        )


def _records_for_file(
    path: Path,
    extension: str,
    *,
    file_password: str | None = None,
    source_mapping_profile: SourceMappingProfile | None = None,
) -> Iterable[tuple[str, dict[object, object]]]:
    if extension == ".xls":
        raise LegacyExcelConversionRequired(
            "Formato XLS legacy non disponibile; salvare una copia come XLSX."
        )
    if extension == ".xlsx":
        return _xlsx_records(path, file_password)
    if extension == ".docx":
        return _docx_records(path, source_mapping_profile)
    if extension == ".pdf":
        return _pdf_records(path, source_mapping_profile)

    raw = path.read_bytes()
    if extension == ".xml":
        return _xml_records(raw, source_mapping_profile)
    text = _decode_text(raw)
    if extension == ".csv":
        return _csv_records(text, ",")
    if extension == ".tsv":
        return _csv_records(text, "\t")
    if extension == ".json":
        return _json_records(text, source_mapping_profile)
    if extension in {".ini", ".cfg"}:
        return _ini_records(text, source_mapping_profile)
    return _key_value_records(text, source_mapping_profile=source_mapping_profile)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _safe_selected_path(root: Path, supplied: str) -> tuple[Path, str]:
    relative = Path(supplied)
    if relative.is_absolute():
        raise ReviewSourceUnavailable(
            "Il file selezionato deve essere relativo alla cartella clienti."
        )
    unresolved_candidate = root / relative
    if unresolved_candidate.is_symlink():
        raise ReviewSourceUnavailable(
            "I collegamenti a file non vengono aperti in revisione."
        )
    candidate = unresolved_candidate.resolve()
    try:
        normalized_relative = candidate.relative_to(root)
    except ValueError as exc:
        raise ReviewSourceUnavailable(
            "Percorso selezionato esterno alla cartella clienti."
        ) from exc
    if not candidate.is_file():
        raise ReviewSourceUnavailable(
            "File selezionato non trovato o non accessibile."
        )
    return candidate, normalized_relative.as_posix()


def _normalized_supplied_path(value: object) -> str:
    return str(value).replace("\\", "/").strip()


def _safe_extension_bucket(value: object) -> str:
    """Return a bounded source-format label without retaining a filename."""

    name = Path(_normalized_supplied_path(value)).name
    extension = ".env" if name.lower() == ".env" else Path(name).suffix.lower()
    if not extension:
        return "(senza estensione)"
    if re.fullmatch(r"\.[a-z0-9]{1,12}", extension):
        return extension
    return "(altro)"


def analyze_files(
    root: str | Path,
    selected_files: Iterable[str],
    *,
    file_passwords: Mapping[str, str] | None = None,
    source_mapping_profile: SourceMappingProfile | Mapping[str, object] | None = None,
) -> ReviewResult:
    root_path = Path(root).expanduser().resolve()
    if not root_path.is_dir():
        raise ReviewError("La cartella clienti non esiste o non è accessibile.")

    supplied_files = list(dict.fromkeys(str(value) for value in selected_files))
    if not supplied_files:
        raise ReviewError("Selezionare almeno un file da revisionare.")
    profile = (
        source_mapping_profile
        if isinstance(source_mapping_profile, SourceMappingProfile)
        else normalize_source_mapping_profile(source_mapping_profile)
    )
    supplied_keys = {_normalized_supplied_path(value) for value in supplied_files}
    password_map: dict[str, str] = {}
    for supplied_path, password in dict(file_passwords or {}).items():
        normalized_path = _normalized_supplied_path(supplied_path)
        if normalized_path not in supplied_keys:
            raise ReviewError("È stata fornita una password per un file non selezionato.")
        if not isinstance(password, str) or not password:
            raise ReviewError("La password di un file Excel selezionato non è valida.")
        if len(password) > MAX_OFFICE_PASSWORD_CHARACTERS:
            raise ReviewError(
                f"La password di un file Excel supera {MAX_OFFICE_PASSWORD_CHARACTERS} caratteri."
            )
        password_map[normalized_path] = password

    candidates: list[CredentialCandidate] = []
    warnings: list[str] = []
    protected_excel_issues: list[ProtectedExcelIssue] = []
    issue_counts: dict[tuple[str, str], int] = {}
    analyzed_files = 0

    def add_issue(reason_code: str, source: object) -> None:
        key = (reason_code, _safe_extension_bucket(source))
        issue_counts[key] = issue_counts.get(key, 0) + 1

    for supplied in supplied_files:
        relative_path = _normalized_supplied_path(supplied)
        try:
            path, relative_path = _safe_selected_path(root_path, supplied)
            extension = path.suffix.lower()
            if not extension and path.name.lower() == ".env":
                extension = ".env"
            if extension not in REVIEWABLE_EXTENSIONS:
                raise UnsupportedReviewFormat(
                    f"Formato {extension or '(senza estensione)'} non revisionabile."
                )
            size = path.stat().st_size
            if size > MAX_FILE_BYTES:
                raise ReviewFileTooLarge(
                    f"File oltre il limite di {MAX_FILE_BYTES // (1024 * 1024)} MB."
                )

            source_hash = _sha256(path)
            relative_parts = Path(relative_path).parts
            client = relative_parts[0] if len(relative_parts) > 1 else ROOT_CLIENT_LABEL
            file_password = password_map.get(relative_path)
            source_password_required = (
                extension == ".xlsx" and _is_ole_compound_file(path)
            )
            file_candidates = 0
            for location, record in _records_for_file(
                path,
                extension,
                file_password=file_password,
                source_mapping_profile=profile,
            ):
                candidate = _make_candidate(
                    record,
                    relative_path=relative_path,
                    source_hash=source_hash,
                    client=client,
                    location=location,
                    source_password_required=source_password_required,
                    source_mapping_profile=profile,
                )
                if candidate is not None:
                    candidates.append(candidate)
                    file_candidates += 1
            analyzed_files += 1
            if file_candidates == 0:
                warnings.append(f"{relative_path}: nessun candidato riconosciuto.")
                add_issue("no_candidate", relative_path)
        except ExcelPasswordRequired:
            protected_excel_issues.append(
                ProtectedExcelIssue(relative_path=relative_path, status="password_required")
            )
            add_issue("password_required", relative_path)
        except ExcelPasswordRejected:
            protected_excel_issues.append(
                ProtectedExcelIssue(relative_path=relative_path, status="password_rejected")
            )
            add_issue("password_rejected", relative_path)
        except ExcelEncryptionReaderUnavailable:
            protected_excel_issues.append(
                ProtectedExcelIssue(relative_path=relative_path, status="reader_unavailable")
            )
            add_issue("reader_unavailable", relative_path)
        except LegacyExcelConversionRequired as exc:
            warnings.append(f"{supplied}: {exc}")
            add_issue("legacy_xls_conversion", relative_path)
        except UnsupportedReviewFormat as exc:
            warnings.append(f"{supplied}: {exc}")
            add_issue("unsupported_review_format", relative_path)
        except ReviewFileTooLarge as exc:
            warnings.append(f"{supplied}: {exc}")
            add_issue("file_too_large", relative_path)
        except ReviewSourceUnavailable as exc:
            warnings.append(f"{supplied}: {exc}")
            add_issue("source_not_available", relative_path)
        except (OSError, ReviewError, ValueError, zipfile.BadZipFile) as exc:
            warnings.append(f"{supplied}: {exc}")
            add_issue("document_unreadable", relative_path)
        except Exception as exc:
            warnings.append(f"{supplied}: documento non analizzabile ({type(exc).__name__}).")
            add_issue("document_unreadable", relative_path)

    ready_count = sum(candidate.status == "ready" for candidate in candidates)
    return ReviewResult(
        root=str(root_path),
        reviewed_at_utc=_utc_now(),
        selected_files=len(supplied_files),
        analyzed_files=analyzed_files,
        candidate_count=len(candidates),
        ready_count=ready_count,
        incomplete_count=len(candidates) - ready_count,
        candidates=candidates,
        warnings=warnings,
        protected_excel_issues=protected_excel_issues,
        issues=[
            ReviewIssueSummary(reason_code, extension, count)
            for (reason_code, extension), count in sorted(issue_counts.items())
        ],
        source_mapping_profile_name=(
            profile.name if profile is not None else "Rilevamento automatico"
        ),
        source_mapping_profile_digest=(profile.digest if profile is not None else ""),
    )


def _read_bounded_json_stdin(maximum_bytes: int, label: str) -> object:
    raw = sys.stdin.buffer.read(maximum_bytes + 1)
    if len(raw) > maximum_bytes:
        raise ReviewError(f"{label} è troppo grande.")
    try:
        return json.loads(raw.decode("utf-8-sig"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ReviewError(f"{label} non contiene JSON valido.") from exc


def _read_secure_review_request() -> tuple[
    list[str], dict[str, str], SourceMappingProfile | None
]:
    document = _read_bounded_json_stdin(
        MAX_SECURE_REVIEW_STDIN_BYTES, "La richiesta di revisione protetta"
    )
    if not isinstance(document, dict):
        raise ReviewError("La richiesta di revisione protetta deve essere un oggetto JSON.")
    files = document.get("files")
    if not isinstance(files, list) or not files or not all(isinstance(item, str) for item in files):
        raise ReviewError("La richiesta protetta non contiene file selezionati validi.")
    password_entries = document.get("file_passwords", [])
    if not isinstance(password_entries, list):
        raise ReviewError("L’elenco delle password Excel non è valido.")
    passwords: dict[str, str] = {}
    for entry in password_entries:
        if not isinstance(entry, dict):
            raise ReviewError("Una password Excel non è valida.")
        relative_path = _normalized_supplied_path(entry.get("relative_path", ""))
        password = entry.get("password")
        if (
            not relative_path
            or relative_path in passwords
            or not isinstance(password, str)
            or not password
            or len(password) > MAX_OFFICE_PASSWORD_CHARACTERS
        ):
            raise ReviewError("Una password Excel non è valida.")
        passwords[relative_path] = password
    profile = normalize_source_mapping_profile(document.get("source_mapping_profile"))
    return files, passwords, profile


def _read_source_mapping_profile_file(path_value: str) -> SourceMappingProfile:
    path = Path(path_value).expanduser().resolve()
    try:
        if not path.is_file() or path.stat().st_size > MAX_SOURCE_MAPPING_PROFILE_BYTES:
            raise ReviewError("Il file del profilo di mappatura non è valido o è troppo grande.")
        document = json.loads(path.read_text(encoding="utf-8-sig"))
    except ReviewError:
        raise
    except (OSError, UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ReviewError("Il file del profilo di mappatura non contiene JSON valido.") from exc
    profile = normalize_source_mapping_profile(document)
    if profile is None:
        raise ReviewError("Il file non contiene un profilo di mappatura.")
    return profile


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Revisione locale e mascherata dei candidati Passbolt.")
    parser.add_argument("--root", help="Cartella principale configurata nell'app.")
    parser.add_argument("--file", action="append", default=[], help="Percorso relativo selezionato; ripetibile.")
    parser.add_argument("--json", action="store_true", help="Stampa il report JSON con segreti mascherati.")
    parser.add_argument(
        "--secure-json",
        action="store_true",
        help="Legge file e password Excel da stdin e restituisce una busta JSON.",
    )
    parser.add_argument(
        "--source-profile",
        help="Profilo JSON locale con mapping esatto dei campi sorgente.",
    )
    parser.add_argument(
        "--profile-check",
        action="store_true",
        help="Valida da stdin un profilo sorgente e restituisce forma canonica e digest.",
    )
    parser.add_argument("--self-test", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.self_test:
        print(
            json.dumps(
                {
                    "version": APP_VERSION,
                    "unlimited_file_selection": True,
                    "unlimited_candidate_collection": True,
                    "single_pass_field_detection": True,
                    "source_mapping_profiles": True,
                    "source_mapping_profile_schema": SOURCE_MAPPING_SCHEMA_VERSION,
                    "reviewable_extensions": len(REVIEWABLE_EXTENSIONS),
                    "excel_password_prompt_supported": True,
                    "secrets_serialized": False,
                }
            )
        )
        return 0
    if args.profile_check:
        try:
            document = _read_bounded_json_stdin(
                MAX_SOURCE_MAPPING_PROFILE_BYTES,
                "Il profilo di mappatura sorgente",
            )
            if isinstance(document, Mapping) and "source_mapping_profile" in document:
                document = document.get("source_mapping_profile")
            profile = normalize_source_mapping_profile(document)
            if profile is None:
                raise ReviewError("Il profilo di mappatura sorgente è mancante.")
            print(
                json.dumps(
                    {"ok": True, "result": {"profile": profile.document()}},
                    ensure_ascii=True,
                    separators=(",", ":"),
                )
            )
            return 0
        except ReviewError as exc:
            print(
                json.dumps(
                    {
                        "ok": False,
                        "error": {
                            "code": "SOURCE_MAPPING_PROFILE_INVALID",
                            "message": str(exc),
                        },
                    },
                    ensure_ascii=True,
                    separators=(",", ":"),
                )
            )
            return 2
    if args.secure_json:
        files: list[str] = []
        passwords: dict[str, str] = {}
        profile: SourceMappingProfile | None = None
        try:
            if not args.root:
                raise ReviewError("La cartella clienti non è configurata.")
            files, passwords, profile = _read_secure_review_request()
            result = analyze_files(
                args.root,
                files,
                file_passwords=passwords,
                source_mapping_profile=profile,
            )
            envelope = {"ok": True, "result": asdict(result)}
            print(json.dumps(envelope, ensure_ascii=True, separators=(",", ":")))
            return 0
        except ReviewError as exc:
            print(
                json.dumps(
                    {
                        "ok": False,
                        "error": {
                            "code": "SECURE_REVIEW_FAILED",
                            "message": str(exc),
                        },
                    },
                    ensure_ascii=True,
                    separators=(",", ":"),
                )
            )
            return 2
        finally:
            passwords.clear()
            files.clear()

    if not args.root or not args.file:
        print("ERRORE: indicare --root e almeno un --file.", file=sys.stderr)
        return 2
    try:
        profile = (
            _read_source_mapping_profile_file(args.source_profile)
            if args.source_profile
            else None
        )
        result = analyze_files(
            args.root, args.file, source_mapping_profile=profile
        )
    except ReviewError as exc:
        print(f"ERRORE: {exc}", file=sys.stderr)
        return 2

    document = asdict(result)
    if args.json:
        # ensure_ascii keeps the stdout contract reliable under Windows
        # PowerShell 5.1 regardless of its active console code page.
        print(json.dumps(document, indent=2, ensure_ascii=True))
    else:
        print(f"File analizzati: {result.analyzed_files}/{result.selected_files}")
        print(f"Candidati: {result.candidate_count}")
        print(f"Pronti: {result.ready_count}")
        print(f"Da completare: {result.incomplete_count}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
